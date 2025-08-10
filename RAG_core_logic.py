import warnings
import os
import logging

# Suppress all warnings
warnings.filterwarnings("ignore")

# Suppress specific torch/transformers warnings
warnings.filterwarnings("ignore", category=FutureWarning, module="torch")
warnings.filterwarnings("ignore", message=".*encoder_attention_mask.*")
warnings.filterwarnings("ignore", message=".*deprecated.*")

# Set environment variables to suppress warnings
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
os.environ['TOKENIZERS_PARALLELISM'] = 'false'

# Configure logging levels
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("torch").setLevel(logging.ERROR)
logging.getLogger("sentence_transformers").setLevel(logging.ERROR)

import asyncio
from typing import Dict, List, Any, Optional, TypedDict
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_anthropic import ChatAnthropic
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph, END, START
from langgraph.graph.message import add_messages
from langgraph.checkpoint.memory import MemorySaver
import operator
from typing_extensions import Annotated
from dotenv import load_dotenv
import json
import PyPDF2
import docx

load_dotenv()

# Initialize LLM with error handling
try:
    llm = ChatAnthropic(
        model="claude-3-haiku-20240307",  
        temperature=0,
        max_tokens=100
    )
except Exception as e:
    print(f"Error initializing LLM: {e}")
    print("Please check your ANTHROPIC_API_KEY in .env file")
    # Don't exit here for Chainlit - let it continue
    llm = None

# Use OpenAI embeddings
try:
    embedding = OpenAIEmbeddings(
        model="text-embedding-3-small",  # or "text-embedding-ada-002" for older model
        openai_api_key=os.getenv("OPENAI_API_KEY")
    )
except Exception as e:
    print(f"Error initializing OpenAI embeddings: {e}")
    print("Please check your OPENAI_API_KEY in .env file")
    embedding = None

vector_store = None

class pdfRAGState(TypedDict):
    messages: Annotated[List[HumanMessage | AIMessage], add_messages]
    original_query: str
    rewritten_query: str
    retrieved_context: str
    final_response: str
    conversation_summary: str  
    message_count: int  

def rewrite_query(state: pdfRAGState) -> pdfRAGState:
    """Step 1: Rewrite the initial query for better retrieval"""
    rewrite_prompt = ChatPromptTemplate.from_template(
        """You are an expert at rewriting user queries to improve information retrieval.

        Original query: {query}

        Rewrite this query to be more specific and effective for information retrieval.
        Consider:
        - Adding relevant keywords
        - Making the intent clearer
        - Expanding abbreviations
        - Adding context that might be implicit

        Return only the rewritten query, nothing else."""
    )

    original_query = state["original_query"]
    
    try:
        if llm:
            chain = rewrite_prompt | llm | StrOutputParser()
            rewritten_query = chain.invoke({"query": original_query})
            state["rewritten_query"] = rewritten_query
        else:
            state["rewritten_query"] = original_query
    except Exception as e:
        print(f"Error in query rewriting: {e}")
        # Fallback to original query if rewriting fails
        state["rewritten_query"] = original_query
    
    return state

def read_docx_with_python_docx(file_path):
    """
    Read .docx files using python-docx library
    Best for modern Word documents (.docx)
    """
    try:
        doc = docx.Document(file_path)
        content = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"
        
        return content.strip()
    except Exception as e:
        print(f"Error reading .docx file: {e}")
        return None

def read_txt_file(file_path):
    """Read plain text files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content.strip()
    except Exception as e:
        print(f"Error reading .txt file: {e}")
        return None

def load_documents(file_path: str):
    """Load documents from PDF, JSON, DOCX, or TXT file into the vector store"""
    global vector_store

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    if not embedding:
        raise ValueError("Embeddings not initialized. Please check your OpenAI API key.")

    try:
        content = ""
        
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Check if text was extracted
                        content += page_text
                        
        elif file_path.endswith('.json'):
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                content = json.dumps(data, indent=2)
                
        elif file_path.endswith('.docx'):
            content = read_docx_with_python_docx(file_path)
            if not content:
                raise ValueError("Could not extract content from DOCX file")
                
        elif file_path.endswith('.txt'):
            content = read_txt_file(file_path)
            if not content:
                raise ValueError("Could not extract content from TXT file")
                
        else:
            raise ValueError("Supported formats: PDF, JSON, DOCX, TXT files")

        # Check if content was extracted successfully
        if not content or len(content.strip()) == 0:
            raise ValueError("No content could be extracted from the file")

        print(f"Extracted {len(content)} characters from {file_path}")

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        splits = text_splitter.split_text(content)

        # Check if splits were created successfully
        if not splits or len(splits) == 0:
            raise ValueError("No text chunks were created from the content")

        print(f"Created {len(splits)} text chunks")

        # Create embeddings and vector store
        vector_store = FAISS.from_texts(splits, embedding)
        print(f"Successfully loaded {len(splits)} document chunks into vector store")

    except Exception as e:
        print(f"Error loading documents: {str(e)}")
        raise

def load_multiple_documents(file_paths: List[str]):
    """Load multiple documents into the same vector store"""
    global vector_store
    
    if not embedding:
        raise ValueError("Embeddings not initialized. Please check your OpenAI API key.")
    
    all_texts = []
    all_metadatas = []
    
    for file_path in file_paths:
        try:
            # Extract content based on file type
            content = ""
            file_name = os.path.basename(file_path)
            
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text
                            
            elif file_path.endswith('.json'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    content = json.dumps(data, indent=2)
                    
            elif file_path.endswith('.docx'):
                content = read_docx_with_python_docx(file_path)
                
            elif file_path.endswith('.txt'):
                content = read_txt_file(file_path)
            
            if content:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len,
                    separators=["\n\n", "\n", " ", ""]
                )
                splits = text_splitter.split_text(content)
                
                # Add metadata for each chunk
                for split in splits:
                    all_texts.append(split)
                    all_metadatas.append({"source": file_name})
                
                print(f"Processed {file_name}: {len(splits)} chunks")
                
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            continue
    
    if all_texts:
        vector_store = FAISS.from_texts(all_texts, embedding, metadatas=all_metadatas)
        print(f"Successfully loaded {len(all_texts)} total chunks from {len(file_paths)} documents")
    else:
        raise ValueError("No content could be extracted from any of the files")

def retrieve_context(state: pdfRAGState) -> pdfRAGState:
    """Step: Retrieve context from vector database"""
    query = state["rewritten_query"]

    try:
        if vector_store:
            docs = vector_store.similarity_search(query, k=3)
            if not docs:
                context = "No relevant documents found in the vector database."
            else:
                context = "\n\n".join([doc.page_content for doc in docs])
        else:
            context = "Vector database not initialized. Please load documents first."

    except Exception as e:
        context = f"Error retrieving context: {str(e)}"
        print(f"Error retrieving context: {str(e)}")

    state["retrieved_context"] = context
    return state

def detect_memory_needed_query(query: str) -> bool:
    """
    Intelligently detects if user is asking about previous conversation
    Only includes memory context when actually needed
    """
    memory_indicators = [
        # Direct references
        "what did i ask", "what did i say", "previously", "before", "earlier", 
        "last question", "last time", "you said", "you mentioned", "you told me",
        # Conversational references  
        "that pdf", "the document we discussed", "from before", "mentioned earlier",
        # Contextual references
        "continue", "also", "in addition to", "regarding that", "about what we"
    ]

    query_lower = query.lower()
    return any(indicator in query_lower for indicator in memory_indicators)

def generate_contextual_pdf_response(state: pdfRAGState) -> pdfRAGState:
    """
    Generates intelligent responses with cost-optimized memory usage
    Only includes conversation history when user specifically asks for it
    """
    if not llm:
        state["final_response"] = "LLM not initialized. Please check your Anthropic API key."
        return state
        
    vector_context = state.get("retrieved_context", "")
    current_query = state["original_query"]

    # Smart memory inclusion - only when needed
    memory_context = ""
    if detect_memory_needed_query(current_query):
        # Include summary of older conversation
        if state.get("conversation_summary"):
            memory_context += f"{state['conversation_summary']}\n\n"

        # Add recent exchanges (last 3 messages max)
        recent_messages = state["messages"][-3:] if len(state["messages"]) >= 4 else state["messages"]
        recent_context = []

        for msg in recent_messages:
            if isinstance(msg, HumanMessage):
                recent_context.append(f"User: {msg.content}")
            elif isinstance(msg, AIMessage):
                # Truncate long AI responses
                ai_content = msg.content[:150] + "..." if len(msg.content) > 150 else msg.content
                recent_context.append(f"Assistant: {ai_content}")

        if recent_context:
            memory_context += "Recent conversation:\n" + "\n".join(recent_context)

    # Optimized prompt template
    if memory_context:
        response_prompt = ChatPromptTemplate.from_template(
            """You are a document analysis assistant specializing in RAG responses.

            {memory_context}

            Current Query: {original_query}
            Enhanced Query: {rewritten_query}

            Document Context from Knowledge Base:
            {context}

            Provide a helpful, accurate response. Reference previous conversation only when relevant.
            Stay focused on the document content and analysis. Be comprehensive but concise.

            Response:"""
        )
    else:
        response_prompt = ChatPromptTemplate.from_template(
            """You are a document analysis assistant specializing in RAG responses.

            Query: {original_query}
            Enhanced Query: {rewritten_query}

            Document Context from Knowledge Base:
            {context}

            Provide a helpful, accurate response about the document content.
            Stay focused on document analysis. Be comprehensive but concise.

            Response:"""
        )

    try:
        cost_optimized_llm = ChatAnthropic(
            model="claude-3-haiku-20240307", 
            temperature=0,
            max_tokens=500  # Increased for better responses
        )

        chain = response_prompt | cost_optimized_llm | StrOutputParser()

        if memory_context:
            response = chain.invoke({
                "memory_context": memory_context,
                "original_query": current_query,
                "rewritten_query": state["rewritten_query"],
                "context": vector_context
            })
        else:
            response = chain.invoke({
                "original_query": current_query,
                "rewritten_query": state["rewritten_query"],
                "context": vector_context
            })

        state["final_response"] = response
        
    except Exception as e:
        print(f"Error generating response: {e}")
        state["final_response"] = f"I encountered an error processing your query: {str(e)}. Please try again."
    
    return state

def manage_conversation_memory(state: pdfRAGState, window_size: int = 4) -> pdfRAGState:
    """
    Maintains sliding window of recent messages + summary of older ones
    Keeps costs predictable regardless of conversation length
    """
    messages = state.get("messages", [])

    if len(messages) > window_size:
        # Keep recent messages
        recent_messages = messages[-window_size:]

        # Create summary of older messages if not exists
        if not state.get("conversation_summary"):
            older_messages = messages[:-window_size]
            summary_parts = []

            for i in range(0, len(older_messages), 2):  # Process in pairs
                if i < len(older_messages):
                    if isinstance(older_messages[i], HumanMessage):
                        user_msg = older_messages[i].content[:80] + "..." if len(older_messages[i].content) > 80 else older_messages[i].content
                        summary_parts.append(f"User asked: {user_msg}")

                if i + 1 < len(older_messages):
                    if isinstance(older_messages[i + 1], AIMessage):
                        ai_msg = older_messages[i + 1].content[:80] + "..." if len(older_messages[i + 1].content) > 80 else older_messages[i + 1].content
                        summary_parts.append(f"AI responded: {ai_msg}")

            state["conversation_summary"] = "Previous conversation:\n" + "\n".join(summary_parts)

        state["messages"] = recent_messages

    state["message_count"] = len(state.get("messages", []))
    return state

def create_optimized_pdf_rag_graph():
    """
    Creates cost-optimized RAG graph with intelligent memory management
    """
    checkpointer = MemorySaver()
    graph = StateGraph(pdfRAGState)

    graph.add_node("memory_optimizer", manage_conversation_memory)
    graph.add_node("query_enhancer", rewrite_query)  
    graph.add_node("pdf_context_retriever", retrieve_context) 
    graph.add_node("intelligent_responder", generate_contextual_pdf_response)

    graph.add_edge(START, "memory_optimizer")
    graph.add_edge("memory_optimizer", "query_enhancer")
    graph.add_edge("query_enhancer", "pdf_context_retriever")
    graph.add_edge("pdf_context_retriever", "intelligent_responder")
    graph.add_edge("intelligent_responder", END)

    return graph.compile(checkpointer=checkpointer)

def manage_thread_lifecycle(thread_id: str, message_count: int, reset_threshold: int = 30) -> str:
    """
    Automatically manages thread lifecycle to prevent infinite cost growth
    Resets thread after threshold to maintain predictable costs
    """
    if message_count > reset_threshold:
        reset_number = message_count // reset_threshold
        return f"{thread_id}_session_{reset_number}"
    return thread_id

optimized_rag_model = None

def initialize_optimized_model():
    """Initialize the cost-optimized model once and reuse it"""
    global optimized_rag_model
    if optimized_rag_model is None:
        optimized_rag_model = create_optimized_pdf_rag_graph()
    return optimized_rag_model

def process_pdf_query_optimized(question: str, thread_id: str = 'pdf_chat_1') -> dict:
    """
    Cost-optimized PDF query processing with intelligent memory management
    Main function for processing user queries about uploaded PDFs
    """
    # Get or create optimized model
    model = initialize_optimized_model()

    config = {'configurable': {'thread_id': thread_id}}

    try:
        current_state = model.get_state(config)
        message_count = len(current_state.values.get('messages', [])) if current_state and current_state.values else 0
    except:
        message_count = 0

    effective_thread_id = manage_thread_lifecycle(thread_id, message_count)
    config = {'configurable': {'thread_id': effective_thread_id}}

    initial_state = pdfRAGState(
        messages=[HumanMessage(content=question)],
        original_query=question,
        rewritten_query="",
        retrieved_context="",
        final_response="",
        conversation_summary="",
        message_count=0
    )

    try:
        result = model.invoke(initial_state, config=config)

        # Add AI response to memory only if it might be referenced later
        if detect_memory_needed_query(question) or len(question) > 50:  # Longer queries might need follow-up
            result["messages"].append(AIMessage(content=result["final_response"]))

        return {
            "response": result["final_response"],
            "context": result.get("retrieved_context", ""),
            "thread_id": effective_thread_id,
            "memory_used": len(result.get("conversation_summary", "")) > 0
        }
    except Exception as e:
        print(f"Error processing query: {e}")
        return {
            "response": f"I encountered an error processing your query. Please try again.",
            "context": "",
            "thread_id": effective_thread_id,
            "memory_used": False
        }

# Remove the main execution block for Chainlit compatibility
# The following code will only run when this file is executed directly, not when imported
if __name__ == "__main__":
    print("🚀 Initializing Document RAG System...")
    
    # Check if document file exists before loading
    doc_path = "file.docx"  # Change this to your DOCX file path
    
    if not os.path.exists(doc_path):
        print(f"❌ Error: Document file not found at {doc_path}")
        print("Please check the file path and try again.")
        print("Supported formats: .pdf, .docx, .json, .txt")
        exit(1)

    try:
        # Suppress warnings during loading
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            load_documents(doc_path)
    except Exception as e:
        print(f"❌ Error loading documents: {e}")
        exit(1)

    print("\n🎥 Smart Document RAG System Ready!")
    print("=" * 50)
    print("Ask me anything about your uploaded document (PDF, DOCX, JSON, TXT).")
    print("Type 'quit' to exit.")
    print("=" * 50)

    thread_id = 'pdf_analysis_session_1'
    message_counter = 0

    while True:
        try:
            user_message = input("\n💬 Your question: ")
            if user_message.strip().lower() in ["q", "exit", "stop", "bye", "goodbye", "quit", "close"]:
                break

            message_counter += 1

            # Clean processing with suppressed warnings
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                result = process_pdf_query_optimized(user_message, thread_id)

            # Clean output format
            print(f"\n🤖 {result['response']}")

            # Optional: Show memory usage (remove if you don't want this)
            if result.get('memory_used'):
                print("   💭 (Referenced previous conversation)")

            thread_id = result['thread_id']

        except KeyboardInterrupt:
            print("\n\nSession interrupted by user.")
            break
        except Exception as e:
            print(f"\n❌ Something went wrong. Please try again.")
            print(f"Debug: {str(e)}")

    print("\n👋 Thanks for using Document RAG System!")