import warnings
import os
import logging
import tempfile
import streamlit as st

# Suppress all warnings
warnings.filterwarnings("ignore")
warnings.filterwarnings("ignore", category=FutureWarning, module="torch")
warnings.filterwarnings("ignore", message=".*encoder_attention_mask.*")
warnings.filterwarnings("ignore", message=".*deprecated.*")

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
os.environ['TOKENIZERS_PARALLELISM'] = 'false'

logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("torch").setLevel(logging.ERROR)
logging.getLogger("sentence_transformers").setLevel(logging.ERROR)

import asyncio
from typing import Dict, List, Any, Optional, TypedDict
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_anthropic import ChatAnthropic
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph, END, START
from langgraph.graph.message import add_messages
from langgraph.checkpoint.memory import MemorySaver
import operator
from typing_extensions import Annotated
from dotenv import load_dotenv
import json
import PyPDF2
import docx
import io

load_dotenv()

# Page configuration
st.set_page_config(
    page_title="🚀 Document RAG Assistant",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "documents_loaded" not in st.session_state:
    st.session_state.documents_loaded = False
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "rag_model" not in st.session_state:
    st.session_state.rag_model = None
if "thread_id" not in st.session_state:
    st.session_state.thread_id = "streamlit_session_1"

# Initialize LLM and embeddings
@st.cache_resource
def initialize_llm():
    try:
        llm = ChatAnthropic(
            model="claude-3-haiku-20240307",  
            temperature=0,
            max_tokens=100
        )
        return llm
    except Exception as e:
        st.error(f"Error initializing LLM: {e}")
        st.error("Please check your ANTHROPIC_API_KEY in .env file or Streamlit secrets")
        return None

@st.cache_resource
def initialize_embeddings():
    try:
        embedding = OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
        )
        return embedding
    except Exception as e:
        st.error(f"Error initializing OpenAI embeddings: {e}")
        st.error("Please check your OPENAI_API_KEY in .env file or Streamlit secrets")
        return None

# RAG State definition
class pdfRAGState(TypedDict):
    messages: Annotated[List[HumanMessage | AIMessage], add_messages]
    original_query: str
    rewritten_query: str
    retrieved_context: str
    final_response: str
    conversation_summary: str  
    message_count: int  

def rewrite_query(state: pdfRAGState) -> pdfRAGState:
    """Step 1: Rewrite the initial query for better retrieval"""
    llm = initialize_llm()
    
    rewrite_prompt = ChatPromptTemplate.from_template(
        """You are an expert at rewriting user queries to improve information retrieval.

        Original query: {query}

        Rewrite this query to be more specific and effective for information retrieval.
        Consider:
        - Adding relevant keywords
        - Making the intent clearer
        - Expanding abbreviations
        - Adding context that might be implicit

        Return only the rewritten query, nothing else."""
    )

    original_query = state["original_query"]
    
    try:
        if llm:
            chain = rewrite_prompt | llm | StrOutputParser()
            rewritten_query = chain.invoke({"query": original_query})
            state["rewritten_query"] = rewritten_query
        else:
            state["rewritten_query"] = original_query
    except Exception as e:
        print(f"Error in query rewriting: {e}")
        state["rewritten_query"] = original_query
    
    return state

def read_docx_content(file_content):
    """Read .docx files from file content"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"
        
        return content.strip()
    except Exception as e:
        st.error(f"Error reading .docx file: {e}")
        return None

def read_pdf_content(file_content):
    """Read PDF files from file content"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        content = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                content += page_text
        return content.strip()
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None

def read_txt_content(file_content):
    """Read plain text files from file content"""
    try:
        content = file_content.decode('utf-8')
        return content.strip()
    except Exception as e:
        st.error(f"Error reading .txt file: {e}")
        return None

def read_json_content(file_content):
    """Read JSON files from file content"""
    try:
        data = json.loads(file_content.decode('utf-8'))
        content = json.dumps(data, indent=2)
        return content.strip()
    except Exception as e:
        st.error(f"Error reading JSON file: {e}")
        return None

def load_documents_from_uploaded_files(uploaded_files):
    """Load multiple documents from Streamlit uploaded files"""
    embedding = initialize_embeddings()
    
    if not embedding:
        raise ValueError("Embeddings not initialized. Please check your OpenAI API key.")
    
    all_texts = []
    all_metadatas = []
    processed_files = []
    
    for uploaded_file in uploaded_files:
        try:
            content = ""
            file_name = uploaded_file.name
            file_content = uploaded_file.read()
            
            if file_name.endswith('.pdf'):
                content = read_pdf_content(file_content)
            elif file_name.endswith('.json'):
                content = read_json_content(file_content)
            elif file_name.endswith('.docx'):
                content = read_docx_content(file_content)
            elif file_name.endswith('.txt'):
                content = read_txt_content(file_content)
            
            if content:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len,
                    separators=["\n\n", "\n", " ", ""]
                )
                splits = text_splitter.split_text(content)
                
                for split in splits:
                    all_texts.append(split)
                    all_metadatas.append({"source": file_name})
                
                processed_files.append((file_name, len(splits)))
                
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {e}")
            continue
    
    if all_texts:
        vector_store = FAISS.from_texts(all_texts, embedding, metadatas=all_metadatas)
        return vector_store, processed_files, len(all_texts)
    else:
        raise ValueError("No content could be extracted from any of the files")

def retrieve_context(state: pdfRAGState) -> pdfRAGState:
    """Retrieve context from vector database"""
    query = state["rewritten_query"]

    try:
        if st.session_state.vector_store:
            docs = st.session_state.vector_store.similarity_search(query, k=3)
            if not docs:
                context = "No relevant documents found in the vector database."
            else:
                context = "\n\n".join([doc.page_content for doc in docs])
        else:
            context = "Vector database not initialized. Please upload documents first."

    except Exception as e:
        context = f"Error retrieving context: {str(e)}"
        print(f"Error retrieving context: {str(e)}")

    state["retrieved_context"] = context
    return state

def detect_memory_needed_query(query: str) -> bool:
    """Detect if user is asking about previous conversation"""
    memory_indicators = [
        "what did i ask", "what did i say", "previously", "before", "earlier", 
        "last question", "last time", "you said", "you mentioned", "you told me",
        "that pdf", "the document we discussed", "from before", "mentioned earlier",
        "continue", "also", "in addition to", "regarding that", "about what we"
    ]

    query_lower = query.lower()
    return any(indicator in query_lower for indicator in memory_indicators)

def generate_contextual_pdf_response(state: pdfRAGState) -> pdfRAGState:
    """Generate intelligent responses with cost-optimized memory usage"""
    llm = initialize_llm()
    
    if not llm:
        state["final_response"] = "LLM not initialized. Please check your Anthropic API key."
        return state
        
    vector_context = state.get("retrieved_context", "")
    current_query = state["original_query"]

    # Smart memory inclusion - only when needed
    memory_context = ""
    if detect_memory_needed_query(current_query):
        if state.get("conversation_summary"):
            memory_context += f"{state['conversation_summary']}\n\n"

        recent_messages = state["messages"][-3:] if len(state["messages"]) >= 4 else state["messages"]
        recent_context = []

        for msg in recent_messages:
            if isinstance(msg, HumanMessage):
                recent_context.append(f"User: {msg.content}")
            elif isinstance(msg, AIMessage):
                ai_content = msg.content[:150] + "..." if len(msg.content) > 150 else msg.content
                recent_context.append(f"Assistant: {ai_content}")

        if recent_context:
            memory_context += "Recent conversation:\n" + "\n".join(recent_context)

    # Optimized prompt template
    if memory_context:
        response_prompt = ChatPromptTemplate.from_template(
            """You are a document analysis assistant specializing in RAG responses.

            {memory_context}

            Current Query: {original_query}
            Enhanced Query: {rewritten_query}

            Document Context from Knowledge Base:
            {context}

            Provide a helpful, accurate response. Reference previous conversation only when relevant.
            Stay focused on the document content and analysis. Be comprehensive but concise.

            Response:"""
        )
    else:
        response_prompt = ChatPromptTemplate.from_template(
            """You are a document analysis assistant specializing in RAG responses.

            Query: {original_query}
            Enhanced Query: {rewritten_query}

            Document Context from Knowledge Base:
            {context}

            Provide a helpful, accurate response about the document content.
            Stay focused on document analysis. Be comprehensive but concise.

            Response:"""
        )

    try:
        cost_optimized_llm = ChatAnthropic(
            model="claude-3-haiku-20240307", 
            temperature=0,
            max_tokens=500
        )

        chain = response_prompt | cost_optimized_llm | StrOutputParser()

        if memory_context:
            response = chain.invoke({
                "memory_context": memory_context,
                "original_query": current_query,
                "rewritten_query": state["rewritten_query"],
                "context": vector_context
            })
        else:
            response = chain.invoke({
                "original_query": current_query,
                "rewritten_query": state["rewritten_query"],
                "context": vector_context
            })

        state["final_response"] = response
        
    except Exception as e:
        print(f"Error generating response: {e}")
        state["final_response"] = f"I encountered an error processing your query: {str(e)}. Please try again."
    
    return state

def manage_conversation_memory(state: pdfRAGState, window_size: int = 4) -> pdfRAGState:
    """Maintain sliding window of recent messages + summary of older ones"""
    messages = state.get("messages", [])

    if len(messages) > window_size:
        recent_messages = messages[-window_size:]

        if not state.get("conversation_summary"):
            older_messages = messages[:-window_size]
            summary_parts = []

            for i in range(0, len(older_messages), 2):
                if i < len(older_messages):
                    if isinstance(older_messages[i], HumanMessage):
                        user_msg = older_messages[i].content[:80] + "..." if len(older_messages[i].content) > 80 else older_messages[i].content
                        summary_parts.append(f"User asked: {user_msg}")

                if i + 1 < len(older_messages):
                    if isinstance(older_messages[i + 1], AIMessage):
                        ai_msg = older_messages[i + 1].content[:80] + "..." if len(older_messages[i + 1].content) > 80 else older_messages[i + 1].content
                        summary_parts.append(f"AI responded: {ai_msg}")

            state["conversation_summary"] = "Previous conversation:\n" + "\n".join(summary_parts)

        state["messages"] = recent_messages

    state["message_count"] = len(state.get("messages", []))
    return state

@st.cache_resource
def create_optimized_pdf_rag_graph():
    """Create cost-optimized RAG graph with intelligent memory management"""
    checkpointer = MemorySaver()
    graph = StateGraph(pdfRAGState)

    graph.add_node("memory_optimizer", manage_conversation_memory)
    graph.add_node("query_enhancer", rewrite_query)  
    graph.add_node("pdf_context_retriever", retrieve_context) 
    graph.add_node("intelligent_responder", generate_contextual_pdf_response)

    graph.add_edge(START, "memory_optimizer")
    graph.add_edge("memory_optimizer", "query_enhancer")
    graph.add_edge("query_enhancer", "pdf_context_retriever")
    graph.add_edge("pdf_context_retriever", "intelligent_responder")
    graph.add_edge("intelligent_responder", END)

    return graph.compile(checkpointer=checkpointer)

def manage_thread_lifecycle(thread_id: str, message_count: int, reset_threshold: int = 30) -> str:
    """Automatically manage thread lifecycle to prevent infinite cost growth"""
    if message_count > reset_threshold:
        reset_number = message_count // reset_threshold
        return f"{thread_id}_session_{reset_number}"
    return thread_id

def process_pdf_query_optimized(question: str, thread_id: str = 'streamlit_chat_1') -> dict:
    """Cost-optimized PDF query processing with intelligent memory management"""
    if st.session_state.rag_model is None:
        st.session_state.rag_model = create_optimized_pdf_rag_graph()
    
    model = st.session_state.rag_model
    config = {'configurable': {'thread_id': thread_id}}

    try:
        current_state = model.get_state(config)
        message_count = len(current_state.values.get('messages', [])) if current_state and current_state.values else 0
    except:
        message_count = 0

    effective_thread_id = manage_thread_lifecycle(thread_id, message_count)
    config = {'configurable': {'thread_id': effective_thread_id}}

    initial_state = pdfRAGState(
        messages=[HumanMessage(content=question)],
        original_query=question,
        rewritten_query="",
        retrieved_context="",
        final_response="",
        conversation_summary="",
        message_count=0
    )

    try:
        result = model.invoke(initial_state, config=config)

        if detect_memory_needed_query(question) or len(question) > 50:
            result["messages"].append(AIMessage(content=result["final_response"]))

        return {
            "response": result["final_response"],
            "context": result.get("retrieved_context", ""),
            "thread_id": effective_thread_id,
            "memory_used": len(result.get("conversation_summary", "")) > 0
        }
    except Exception as e:
        print(f"Error processing query: {e}")
        return {
            "response": f"I encountered an error processing your query. Please try again.",
            "context": "",
            "thread_id": effective_thread_id,
            "memory_used": False
        }

# STREAMLIT UI
def main():
    st.title("🚀 Document RAG Assistant")
    st.markdown("Upload documents and ask questions about their content!")

    # Sidebar for document upload
    with st.sidebar:
        st.header("📄 Document Upload")
        
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'txt', 'json'],
            accept_multiple_files=True,
            help="Supported formats: PDF, DOCX, TXT, JSON"
        )
        
        if uploaded_files:
            if st.button("📤 Process Documents"):
                with st.spinner("Processing documents..."):
                    try:
                        vector_store, processed_files, total_chunks = load_documents_from_uploaded_files(uploaded_files)
                        st.session_state.vector_store = vector_store
                        st.session_state.documents_loaded = True
                        
                        st.success(f"✅ Processed {len(processed_files)} documents!")
                        st.write("**Files processed:**")
                        for file_name, chunks in processed_files:
                            st.write(f"• {file_name}: {chunks} chunks")
                        st.write(f"**Total chunks:** {total_chunks}")
                        
                        # Reset thread for new documents
                        st.session_state.thread_id = f"streamlit_session_{hash(''.join([f.name for f in uploaded_files])) % 10000}"
                        
                    except Exception as e:
                        st.error(f"❌ Error processing documents: {str(e)}")
        
        # Document status
        if st.session_state.documents_loaded:
            st.success("✅ Documents ready!")
        else:
            st.warning("⚠️ No documents loaded")
        
        # Clear documents button
        if st.session_state.documents_loaded:
            if st.button("🗑️ Clear Documents"):
                st.session_state.vector_store = None
                st.session_state.documents_loaded = False
                st.session_state.messages = []
                st.rerun()

    # Main chat interface
    if st.session_state.documents_loaded:
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Chat input
        if prompt := st.chat_input("Ask me anything about your documents..."):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            # Generate response
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    result = process_pdf_query_optimized(prompt, st.session_state.thread_id)
                    response = result["response"]
                    st.markdown(response)
                    
                    # Update thread ID
                    st.session_state.thread_id = result["thread_id"]
                    
                    # Show memory usage if enabled
                    if result.get('memory_used'):
                        st.caption("💭 Referenced previous conversation context")

            # Add assistant response to chat history
            st.session_state.messages.append({"role": "assistant", "content": response})

    else:
        st.info("👆 Please upload documents using the sidebar to get started!")
        
        # Example usage
        st.markdown("""
        ### 🎯 How to use:
        1. **Upload Documents** - Use the sidebar to upload PDF, DOCX, TXT, or JSON files
        2. **Process Documents** - Click the "Process Documents" button
        3. **Ask Questions** - Start chatting about your document content
        
        ### 💡 Example questions:
        - "What is this document about?"
        - "Summarize the main points"
        - "What are the key findings?"
        - "Extract information about [specific topic]"
        """)

    # Footer
    st.markdown("---")
    # st.markdown("*Powered by LangChain, Anthropic Claude, and OpenAI Embeddings*")

if __name__ == "__main__":
    main()